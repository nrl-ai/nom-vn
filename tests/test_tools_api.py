"""Tests for /api/tools/* — stateless playground endpoints.

These don't touch spaces / materials / the embedder, so the test client
is configured with the cheapest fakes possible. The HF-backed diacritic
path is excluded — it's gated behind torch + transformers and an actual
model download, so we exercise it via the rule path here and keep the
HF integration covered in `tests/test_data_pipelines.py`.
"""

from __future__ import annotations

import pytest

fastapi = pytest.importorskip("fastapi")
pytest.importorskip("fastapi.testclient")

from fastapi.testclient import TestClient  # noqa: E402

from nom.chat.server import build_app  # noqa: E402
from nom.chat.store import MemoryStore  # noqa: E402
from tests._fakes import FakeEmbedder, FakeLLM  # noqa: E402


@pytest.fixture
def client() -> TestClient:
    store = MemoryStore(embedder=FakeEmbedder(), llm=FakeLLM(response="Đã khôi phục."))
    return TestClient(build_app(store=store))


class TestDiacriticRestore:
    def test_rule_backend_default(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/diacritic/restore",
            json={"text": "toi la nguoi viet nam", "backend": "rule"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["backend"] == "rule"
        # Rule table covers high-freq pronouns/particles.
        assert "tôi" in body["restored"].lower()
        assert "là" in body["restored"].lower()

    def test_missing_text_422(self, client: TestClient) -> None:
        r = client.post("/api/tools/diacritic/restore", json={"text": ""})
        assert r.status_code == 422

    def test_unknown_backend_422(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/diacritic/restore",
            json={"text": "hello", "backend": "magic"},
        )
        assert r.status_code == 422

    def test_llm_backend_uses_provided_llm(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/diacritic/restore",
            json={"text": "toi yeu Viet Nam", "backend": "llm"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["backend"] == "llm"
        # FakeLLM canned response is what we pre-configured above.
        assert body["restored"]

    def test_models_listing(self, client: TestClient) -> None:
        r = client.get("/api/tools/diacritic/models")
        assert r.status_code == 200
        body = r.json()
        assert body["default"] == "nrl-ai/vn-diacritic-vit5-base"
        ids = [m["id"] for m in body["models"]]
        assert "nrl-ai/vn-diacritic-small" in ids


class TestStripDiacritics:
    def test_strip_basic(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/diacritic/strip",
            json={"text": "Hợp đồng số 02/HĐ/2025"},
        )
        assert r.status_code == 200
        assert r.json()["stripped"] == "Hop dong so 02/HD/2025"


class TestTokenize:
    def test_word_list_format(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/tokenize/word",
            json={"text": "Hợp đồng số 02 được lập"},
        )
        assert r.status_code == 200
        body = r.json()
        assert "Hợp đồng" in body["tokens"]
        assert body["n_tokens"] == len(body["tokens"])
        assert body["n_compounds"] >= 1

    def test_word_text_format(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/tokenize/word",
            json={"text": "Hợp đồng số 02", "fmt": "text"},
        )
        assert r.status_code == 200
        assert "Hợp_đồng" in r.json()["text"]

    def test_word_invalid_fmt_422(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/tokenize/word",
            json={"text": "x", "fmt": "json"},
        )
        assert r.status_code == 422

    def test_sentence_tokenize(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/tokenize/sentence",
            json={"text": "Tôi yêu Việt Nam. Bạn có khỏe không? Cảm ơn!"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["n_sentences"] == 3
        assert len(body["sentences"]) == 3


class TestNormalizeDetect:
    def test_normalize_already_nfc(self, client: TestClient) -> None:
        # "Tôi" composed.
        r = client.post(
            "/api/tools/text/normalize",
            json={"text": "Tôi"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["is_nfc"] is True
        assert body["nfc"] == "Tôi"

    def test_normalize_decomposed_input(self, client: TestClient) -> None:
        # "Tôi" with NFD decomposition: o + combining circumflex.
        decomposed = "Tôi"
        r = client.post(
            "/api/tools/text/normalize",
            json={"text": decomposed},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["is_nfc"] is False
        assert body["nfc"] == "Tôi"

    def test_detect_with_diacritics(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/text/detect",
            json={"text": "Hợp đồng số 02"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["is_vietnamese"] is True
        assert body["has_diacritics"] is True

    def test_detect_ascii_vn(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/text/detect",
            json={"text": "Hop dong so hai duoc lap"},
        )
        assert r.status_code == 200
        body = r.json()
        # Stripped form still parses as Vietnamese via the common-word table.
        assert body["has_diacritics"] is False
        assert body["is_vietnamese"] is True

    def test_detect_english(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/text/detect",
            json={"text": "the quick brown fox jumps over the lazy dog"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["has_diacritics"] is False


class TestNLP:
    def test_ner_detects_money_and_date(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/nlp/ner",
            json={"text": "VCB ký hợp đồng 1.500.000 VND ngày 02/05/2026."},
        )
        assert r.status_code == 200
        labels = {s["label"] for s in r.json()["spans"]}
        assert {"MONEY", "DATE", "ORG"}.issubset(labels)

    def test_ner_empty_input_422(self, client: TestClient) -> None:
        r = client.post("/api/tools/nlp/ner", json={"text": ""})
        assert r.status_code == 422

    def test_sentiment_positive(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/nlp/sentiment",
            json={"text": "Sản phẩm này rất tuyệt vời, tôi rất hài lòng."},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["label"] == "positive"
        assert body["score"] > 0.5

    def test_sentiment_negative(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/nlp/sentiment",
            json={"text": "Dịch vụ tệ quá, rất thất vọng."},
        )
        assert r.json()["label"] == "negative"

    def test_language_vn(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/nlp/language",
            json={"text": "Đây là một câu tiếng Việt."},
        )
        assert r.status_code == 200
        assert r.json()["language"] == "vi"

    def test_language_en(self, client: TestClient) -> None:
        r = client.post(
            "/api/tools/nlp/language",
            json={"text": "This is plain English."},
        )
        assert r.json()["language"] == "en"

    def test_language_empty_input_422(self, client: TestClient) -> None:
        r = client.post("/api/tools/nlp/language", json={"text": ""})
        assert r.status_code == 422


class TestTranslate:
    """LLMTranslator path goes through the same FakeLLM the diacritic
    tests use. The fake returns canned JSON; the translator parses
    ``translation`` out and NFC-normalizes it."""

    def _client(self, response: str) -> TestClient:
        store = MemoryStore(embedder=FakeEmbedder(), llm=FakeLLM(response=response))
        return TestClient(build_app(store=store))

    def test_llm_backend_returns_translation(self) -> None:
        client = self._client('{"translation": "Hello world."}')
        r = client.post(
            "/api/tools/translate",
            json={"text": "Xin chào thế giới.", "source": "vi", "target": "en"},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["translation"] == "Hello world."
        assert body["source"] == "vi"
        assert body["target"] == "en"
        assert body["backend"] == "llm"

    def test_missing_text_422(self) -> None:
        client = self._client('{"translation": "X"}')
        r = client.post("/api/tools/translate", json={"text": ""})
        assert r.status_code == 422

    def test_same_source_and_target_422(self) -> None:
        client = self._client('{"translation": "X"}')
        r = client.post(
            "/api/tools/translate",
            json={"text": "Hello", "source": "en", "target": "en"},
        )
        assert r.status_code == 422

    def test_unknown_backend_422(self) -> None:
        client = self._client('{"translation": "X"}')
        r = client.post(
            "/api/tools/translate",
            json={"text": "Hello", "source": "en", "target": "vi", "backend": "magic"},
        )
        assert r.status_code == 422

    def test_unsupported_language_pair_422(self) -> None:
        client = self._client('{"translation": "X"}')
        r = client.post(
            "/api/tools/translate",
            json={"text": "Hello", "source": "fr", "target": "vi"},
        )
        assert r.status_code == 422

    def test_models_listing(self) -> None:
        client = self._client('{"translation": "X"}')
        r = client.get("/api/tools/translate/models")
        assert r.status_code == 200
        body = r.json()
        assert body["default_backend"] == "llm"
        assert "en2vi" in body["directions"]
        assert "vi2en" in body["directions"]
        ids = [m["id"] for m in body["hf_models"]]
        assert "google/madlad400-3b-mt" in ids
        assert "facebook/m2m100_418M" in ids

    def test_falls_back_to_raw_when_llm_returns_non_json(self) -> None:
        client = self._client("Hello world.")
        r = client.post(
            "/api/tools/translate",
            json={"text": "Xin chào thế giới.", "source": "vi", "target": "en"},
        )
        assert r.status_code == 200
        assert r.json()["translation"] == "Hello world."


class TestTranslateFileUpload:
    """File upload → walked .docx → translated .docx flow.

    Skips cleanly when python-docx is not installed (the walker raises
    a clear ImportError that the route would translate to 500). The
    docx extra is in [chat] which the integration tier installs.
    """

    @pytest.fixture
    def docx_bytes(self) -> bytes:
        pytest.importorskip("docx")
        import io

        from docx import Document

        doc = Document()
        doc.add_paragraph("alpha")
        doc.add_paragraph("beta")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _client(self, response: str) -> TestClient:
        store = MemoryStore(embedder=FakeEmbedder(), llm=FakeLLM(response=response))
        return TestClient(build_app(store=store))

    def test_uploads_returns_translated_docx(self, docx_bytes: bytes) -> None:
        import io
        import json as _json

        pytest.importorskip("docx")
        from docx import Document

        client = self._client('{"translation": "ALPHA"}')
        r = client.post(
            "/api/tools/translate/file",
            files={"file": ("source.docx", docx_bytes, "application/octet-stream")},
            data={"source": "vi", "target": "en", "backend": "llm"},
        )
        assert r.status_code == 200
        assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument")
        assert "source.en.docx" in r.headers["content-disposition"]
        stats = _json.loads(r.headers["x-translation-stats"])
        assert stats["paragraphs_translated"] == 2
        out = Document(io.BytesIO(r.content))
        texts = [p.text for p in out.paragraphs if p.text]
        assert texts == ["ALPHA", "ALPHA"]

    def test_rejects_unsupported_format(self) -> None:
        client = self._client('{"translation": "X"}')
        r = client.post(
            "/api/tools/translate/file",
            files={"file": ("source.csv", b"a,b,c\n", "text/csv")},
            data={"source": "vi", "target": "en"},
        )
        assert r.status_code == 422
        assert "unsupported source format" in r.json()["detail"]

    def test_accepts_txt_format(self) -> None:
        client = self._client('{"translation": "Hello"}')
        r = client.post(
            "/api/tools/translate/file",
            files={"file": ("notes.txt", "xin chào".encode(), "text/plain")},
            data={"source": "vi", "target": "en"},
        )
        assert r.status_code == 200
        assert r.headers["content-type"].startswith("text/plain")
        assert "notes.en.txt" in r.headers["content-disposition"]

    def test_rejects_same_source_and_target_file(self, docx_bytes: bytes) -> None:
        client = self._client('{"translation": "X"}')
        r = client.post(
            "/api/tools/translate/file",
            files={"file": ("source.docx", docx_bytes, "application/octet-stream")},
            data={"source": "en", "target": "en"},
        )
        assert r.status_code == 422


class TestConvertFileUpload:
    """Tests for /api/tools/convert/file (PDF / image → DOCX).

    Skips cleanly when Tesseract isn't installed. The image flow uses
    a small in-process PIL fixture; the dispatcher rejection paths are
    deterministic (no OCR).
    """

    @staticmethod
    def _client() -> TestClient:
        store = MemoryStore(embedder=FakeEmbedder(), llm=FakeLLM())
        return TestClient(build_app(store=store))

    @pytest.fixture
    def image_bytes(self) -> bytes:
        pytest.importorskip("PIL")
        import io

        from PIL import Image, ImageDraw, ImageFont

        img = Image.new("L", (640, 200), 255)
        drawer = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 36)
        except OSError:
            font = ImageFont.load_default()
        drawer.text((20, 60), "Hello world", fill=0, font=font)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()

    def test_convert_image_returns_docx(self, image_bytes: bytes) -> None:
        import shutil

        if shutil.which("tesseract") is None:
            pytest.skip("tesseract not installed")
        import io
        import json as _json

        pytest.importorskip("docx")
        from docx import Document

        client = self._client()
        r = client.post(
            "/api/tools/convert/file",
            files={"file": ("scan.png", image_bytes, "image/png")},
            data={"ocr_language": "eng"},
        )
        assert r.status_code == 200
        assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument")
        assert "scan.docx" in r.headers["content-disposition"]
        stats = _json.loads(r.headers["x-convert-stats"])
        assert stats["pages_ocred"] == 1
        assert stats["chars_out"] > 0

        out = Document(io.BytesIO(r.content))
        full = "\n".join(p.text for p in out.paragraphs)
        assert "hello" in full.lower()

    def test_rejects_unsupported_extension(self) -> None:
        client = self._client()
        r = client.post(
            "/api/tools/convert/file",
            files={"file": ("data.csv", b"a,b,c\n", "text/csv")},
            data={"ocr_language": "eng"},
        )
        assert r.status_code == 422
        assert "unsupported source format" in r.json()["detail"]

    def test_rejects_unknown_image_extension(self) -> None:
        client = self._client()
        r = client.post(
            "/api/tools/convert/file",
            files={"file": ("data.xyz", b"\x00\x00", "application/octet-stream")},
            data={"ocr_language": "eng"},
        )
        assert r.status_code == 422
