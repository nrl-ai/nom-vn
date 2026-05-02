"""Unit tests for nom.translate.formats.docx.translate_docx.

Uses python-docx to generate fixtures in-process so no binary fixtures
end up in the repo. Skips cleanly if python-docx is absent.
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

from docx import Document

from nom.translate.formats import DocxTranslationStats, translate_docx


class _Reverse:
    """Identity-shape Translator that reverses each unit. Trivial to
    assert against; covers the round-trip without involving a real LLM."""

    name = "reverse"
    source_lang = "en"
    target_lang = "vi"

    def translate(self, text: str, *, hint: str | None = None) -> str:
        return text[::-1]


class _RaisesOnSecond:
    """Translates the first paragraph, raises on the second. Used to
    verify per-paragraph failure isolation."""

    name = "raises-on-second"
    source_lang = "en"
    target_lang = "vi"

    def __init__(self) -> None:
        self.calls = 0

    def translate(self, text: str, *, hint: str | None = None) -> str:
        self.calls += 1
        if self.calls == 2:
            raise RuntimeError("simulated translator failure")
        return text.upper()


def _save(doc: Document, path: Path) -> None:  # type: ignore[valid-type]
    doc.save(str(path))


def test_translates_body_paragraphs(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("Hello world.")
    doc.add_paragraph("Second paragraph.")
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _Reverse())

    assert isinstance(stats, DocxTranslationStats)
    assert stats.paragraphs_translated == 2
    assert stats.paragraphs_failed == 0
    out = Document(str(dst))
    texts = [p.text for p in out.paragraphs if p.text]
    assert texts == [".dlrow olleH", ".hpargarap dnoceS"]


def test_skips_whitespace_only_paragraphs(tmp_path: Path) -> None:
    """Whitespace-only paragraphs are visited and counted as skipped.
    Truly-empty paragraphs (typical of python-docx default header /
    footer regions) are silently passed through, since counting them
    would clutter the user-visible stats with structural noise."""
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("   ")
    doc.add_paragraph("real content")
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _Reverse())

    assert stats.paragraphs_translated == 1
    assert stats.paragraphs_skipped == 1


def test_translates_table_cells(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "alpha"
    table.cell(0, 1).text = "beta"
    table.cell(1, 0).text = "gamma"
    table.cell(1, 1).text = "delta"
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _Reverse())

    assert stats.paragraphs_translated == 4
    out = Document(str(dst))
    texts = [out.tables[0].cell(r, c).text for r in range(2) for c in range(2)]
    assert texts == ["ahpla", "ateb", "ammag", "atled"]


def test_handles_merged_cells_without_double_translating(tmp_path: Path) -> None:
    """When cells are merged, python-docx returns the same paragraph
    via multiple cell handles — the walker must dedup by paragraph
    identity to avoid translating the same text twice."""
    src = tmp_path / "src.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "merged content"
    table.cell(0, 0).merge(table.cell(0, 1))
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _Reverse())

    assert stats.paragraphs_translated == 1


def test_failure_in_one_paragraph_leaves_others_intact(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph("second")
    doc.add_paragraph("third")
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _RaisesOnSecond())

    assert stats.paragraphs_translated == 2
    assert stats.paragraphs_failed == 1
    out = Document(str(dst))
    texts = [p.text for p in out.paragraphs if p.text]
    assert texts == ["FIRST", "second", "THIRD"]


def test_collapses_runs_into_first_with_remaining_cleared(tmp_path: Path) -> None:
    """v0 strategy: all runs concatenated, translation written into
    first run, others cleared. Confirms run-redistribution behavior."""
    src = tmp_path / "src.docx"
    doc = Document()
    para = doc.add_paragraph("")
    para.add_run("Hello ")
    para.add_run("brave ").bold = True
    para.add_run("world.")
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    translate_docx(src, dst, _Reverse())

    out = Document(str(dst))
    runs = list(out.paragraphs[0].runs)
    assert len(runs) == 3
    assert runs[0].text == ".dlrow evarb olleH"
    assert runs[1].text == ""
    assert runs[2].text == ""


def test_chars_in_out_counts_track_translation_size(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("abc")
    doc.add_paragraph("hello world")
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _Reverse())

    assert stats.chars_in == len("abc") + len("hello world")
    assert stats.chars_out == stats.chars_in  # reverse() preserves length


def test_missing_source_raises_filenotfound(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        translate_docx(tmp_path / "nope.docx", tmp_path / "out.docx", _Reverse())


def test_creates_destination_parent_directory(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("content")
    _save(doc, src)

    dst = tmp_path / "nested" / "subdir" / "out.docx"
    translate_docx(src, dst, _Reverse())
    assert dst.exists()


def test_walks_header_and_footer_paragraphs(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("body text")
    section = doc.sections[0]
    section.header.paragraphs[0].text = "header text"
    section.footer.paragraphs[0].text = "footer text"
    _save(doc, src)

    dst = tmp_path / "dst.docx"
    stats = translate_docx(src, dst, _Reverse())

    assert stats.paragraphs_translated >= 3  # body + header + footer
    out = Document(str(dst))
    assert "txet ydob" in [p.text for p in out.paragraphs]
    assert out.sections[0].header.paragraphs[0].text == "txet redaeh"
    assert out.sections[0].footer.paragraphs[0].text == "txet retoof"
