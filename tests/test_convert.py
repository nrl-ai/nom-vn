"""Tests for nom.convert.{image_to_docx,pdf_to_docx,convert_to_docx}.

Generates fixtures in-process via PIL — no binary fixtures committed.
Skips cleanly when Tesseract isn't installed.
"""

from __future__ import annotations

import shutil
from pathlib import Path

import pytest

pytest.importorskip("docx")
pytest.importorskip("PIL")
pytest.importorskip("pytesseract")

from docx import Document
from PIL import Image, ImageDraw, ImageFont

from nom.convert import (
    ConversionStats,
    convert_to_docx,
    image_to_docx,
)

_TESSERACT_BIN = shutil.which("tesseract")
needs_tesseract = pytest.mark.skipif(
    _TESSERACT_BIN is None,
    reason="tesseract binary not installed",
)


def _make_text_image(
    path: Path, text: str = "Hello world", *, size: tuple[int, int] = (640, 200)
) -> None:
    """Render a grayscale PNG with `text` for OCR test input.

    Uses the default PIL bitmap font so we don't depend on a specific
    truetype file being present. Intentionally large+padded — Tesseract
    needs reasonable letter sizing (~24px+) for reliable recognition.
    """
    img = Image.new("L", size, 255)
    drawer = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 36)
    except OSError:
        font = ImageFont.load_default()
    drawer.text((20, 60), text, fill=0, font=font)
    img.save(str(path))


@needs_tesseract
def test_image_to_docx_basic_round_trip(tmp_path: Path) -> None:
    src = tmp_path / "src.png"
    dst = tmp_path / "out.docx"
    _make_text_image(src, "Hello world")

    stats = image_to_docx(src, dst, ocr_language="eng")

    assert isinstance(stats, ConversionStats)
    assert stats.n_pages == 1
    assert stats.pages_ocred == 1
    assert stats.pages_text_extracted == 0
    assert stats.chars_out > 0
    assert dst.exists()

    doc = Document(str(dst))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "hello" in full_text.lower()


@needs_tesseract
def test_image_to_docx_blank_image_writes_empty_doc(tmp_path: Path) -> None:
    src = tmp_path / "blank.png"
    dst = tmp_path / "out.docx"
    Image.new("L", (200, 200), 255).save(str(src))

    stats = image_to_docx(src, dst, ocr_language="eng")
    assert stats.chars_out == 0
    assert dst.exists()
    doc = Document(str(dst))
    assert len(doc.paragraphs) >= 1


def test_image_to_docx_missing_source_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        image_to_docx(tmp_path / "nope.png", tmp_path / "out.docx")


def test_convert_dispatcher_rejects_unknown_extension(tmp_path: Path) -> None:
    src = tmp_path / "data.csv"
    src.write_text("a,b,c\n", encoding="utf-8")
    with pytest.raises(ValueError, match="unsupported source format"):
        convert_to_docx(src, tmp_path / "out.docx")


@needs_tesseract
def test_convert_dispatcher_routes_to_image_handler(tmp_path: Path) -> None:
    src = tmp_path / "src.jpg"
    dst = tmp_path / "out.docx"
    _make_text_image(src, "Hello world")
    stats = convert_to_docx(src, dst, ocr_language="eng")
    assert stats.pages_ocred == 1
    assert dst.exists()


def test_convert_dispatcher_creates_destination_parent(tmp_path: Path) -> None:
    """Image converter creates intermediate directories under dst."""
    pytest.importorskip("pytesseract")
    if _TESSERACT_BIN is None:
        pytest.skip("tesseract not installed")
    src = tmp_path / "src.png"
    dst = tmp_path / "nested" / "deep" / "out.docx"
    _make_text_image(src, "Hello")
    convert_to_docx(src, dst, ocr_language="eng")
    assert dst.exists()


# ---------------------------------------------------------------------------
# PDF → DOCX paths
# ---------------------------------------------------------------------------


def _make_text_layer_pdf(path: Path, body: str) -> None:
    """Create a one-page PDF with a real text layer via reportlab."""
    rl = pytest.importorskip("reportlab.pdfgen.canvas")
    canvas = rl.Canvas(str(path))
    text_obj = canvas.beginText(72, 720)
    for line in body.splitlines():
        text_obj.textLine(line)
    canvas.drawText(text_obj)
    canvas.showPage()
    canvas.save()


def test_pdf_text_layer_extracted_directly(tmp_path: Path) -> None:
    """Born-digital PDF (text layer present) → DOCX without OCR."""
    pytest.importorskip("reportlab")
    pytest.importorskip("pdfplumber")
    pytest.importorskip("pypdfium2")

    from nom.convert import pdf_to_docx

    src = tmp_path / "src.pdf"
    dst = tmp_path / "out.docx"
    _make_text_layer_pdf(src, "Hello world\nThis is a born-digital PDF.")

    stats = pdf_to_docx(src, dst, ocr_language="eng")

    assert isinstance(stats, ConversionStats)
    assert stats.n_pages == 1
    assert stats.pages_text_extracted == 1
    assert stats.pages_ocred == 0
    assert dst.exists()

    out = Document(str(dst))
    full = "\n".join(p.text for p in out.paragraphs)
    assert "born-digital" in full.lower()


def test_pdf_dispatcher_routes_to_pdf_handler(tmp_path: Path) -> None:
    pytest.importorskip("reportlab")
    pytest.importorskip("pdfplumber")
    pytest.importorskip("pypdfium2")

    src = tmp_path / "doc.pdf"
    dst = tmp_path / "doc.docx"
    _make_text_layer_pdf(
        src,
        "dispatcher test — long enough text to pass the 32-char text-layer threshold.",
    )

    stats = convert_to_docx(src, dst, ocr_language="eng")
    assert stats.n_pages == 1
    assert stats.pages_text_extracted == 1


def test_pdf_missing_source_raises(tmp_path: Path) -> None:
    pytest.importorskip("pdfplumber")
    pytest.importorskip("pypdfium2")

    from nom.convert import pdf_to_docx

    with pytest.raises(FileNotFoundError):
        pdf_to_docx(tmp_path / "nope.pdf", tmp_path / "out.docx")


@needs_tesseract
def test_pdf_with_no_text_layer_falls_back_to_ocr(tmp_path: Path) -> None:
    """Image-only PDF (rendered from a PIL image) — pdfplumber yields
    nothing, so the walker must OCR every page."""
    pytest.importorskip("pdfplumber")
    pytest.importorskip("pypdfium2")
    pytest.importorskip("reportlab")

    from reportlab.pdfgen import canvas as rl_canvas

    from nom.convert import pdf_to_docx

    src = tmp_path / "scan.pdf"
    dst = tmp_path / "scan.docx"

    _make_text_image(tmp_path / "page.png", "Hello world")
    pdf_canvas = rl_canvas.Canvas(str(src), pagesize=(640, 200))
    pdf_canvas.drawImage(str(tmp_path / "page.png"), 0, 0, 640, 200)
    pdf_canvas.showPage()
    pdf_canvas.save()

    stats = pdf_to_docx(src, dst, ocr_language="eng")
    assert stats.pages_text_extracted == 0
    assert stats.pages_ocred == 1
    assert dst.exists()
