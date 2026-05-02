"""Tests for nom.translate._protect — tag-protection placeholder logic."""

from __future__ import annotations

import pytest

from nom.translate._protect import (
    TagProtectResult,
    detect_length_warning,
    translate_with_tag_protection,
)


class _PreservingMockTranslator:
    """Translator that uppercases text but preserves placeholders verbatim."""

    name = "mock-preserving"
    source_lang = "en"
    target_lang = "vi"

    def translate(self, text: str, *, hint: str | None = None) -> str:
        # Upper-case everything, but keep the ⟦N⟧ markers untouched.
        out: list[str] = []
        i = 0
        while i < len(text):
            if text[i] == "⟦":
                end = text.index("⟧", i)
                out.append(text[i : end + 1])
                i = end + 1
            else:
                out.append(text[i].upper())
                i += 1
        return "".join(out)


class _DroppingTranslator:
    """Translator that strips all placeholders — exercises the fallback."""

    name = "mock-dropping"
    source_lang = "en"
    target_lang = "vi"

    def translate(self, text: str, *, hint: str | None = None) -> str:
        out: list[str] = []
        i = 0
        while i < len(text):
            if text[i] == "⟦":
                end = text.index("⟧", i)
                i = end + 1
            else:
                out.append(text[i])
                i += 1
        return "".join(out).upper()


class _ReorderingTranslator:
    """Translator that re-orders placeholders — also fallback territory."""

    name = "mock-reorder"
    source_lang = "en"
    target_lang = "vi"

    def translate(self, text: str, *, hint: str | None = None) -> str:
        # Swap ⟦0⟧ and ⟦1⟧ in the output to force out-of-order detection.
        return text.replace("⟦0⟧", "⟦TMP⟧").replace("⟦1⟧", "⟦0⟧").replace("⟦TMP⟧", "⟦1⟧")


def test_single_run_translates_without_placeholders() -> None:
    result = translate_with_tag_protection(["hello"], _PreservingMockTranslator())
    assert result.protected is True
    assert result.run_texts == ["HELLO"]
    assert result.fallback_text == "HELLO"


def test_empty_run_list_returns_empty_protected_result() -> None:
    result = translate_with_tag_protection([], _PreservingMockTranslator())
    assert result.protected is True
    assert result.run_texts == []


def test_multi_run_preserved_when_translator_keeps_placeholders() -> None:
    runs = ["The ", "contract", " is ", "void", "."]
    result = translate_with_tag_protection(runs, _PreservingMockTranslator())
    assert result.protected is True
    assert result.run_texts == ["THE ", "CONTRACT", " IS ", "VOID", "."]


def test_falls_back_when_translator_drops_placeholders() -> None:
    runs = ["The ", "contract", " is ", "void", "."]
    result = translate_with_tag_protection(runs, _DroppingTranslator())
    assert result.protected is False
    # Fallback text is the placeholder-stripped translation
    assert result.fallback_text == "THE CONTRACT IS VOID."


def test_falls_back_when_placeholders_out_of_order() -> None:
    runs = ["A", "B"]
    result = translate_with_tag_protection(runs, _ReorderingTranslator())
    assert result.protected is False


def test_detect_length_warning_under_threshold() -> None:
    assert detect_length_warning(100, 110) is None


def test_detect_length_warning_over_threshold() -> None:
    msg = detect_length_warning(100, 200)
    assert msg is not None
    assert "2.0×" in msg


def test_detect_length_warning_zero_source() -> None:
    assert detect_length_warning(0, 50) is None


def test_run_preserving_docx_round_trip(tmp_path) -> None:
    """End-to-end: write a docx with mixed-style runs, translate with
    preserve_runs=True, verify each original run got its translated
    text with the same boundaries."""
    pytest.importorskip("docx")
    from docx import Document

    from nom.translate.formats import translate_docx

    src = tmp_path / "src.docx"
    dst = tmp_path / "dst.docx"
    doc = Document()
    para = doc.add_paragraph("")
    para.add_run("The ")
    para.add_run("contract").bold = True
    para.add_run(" is ")
    para.add_run("void").italic = True
    para.add_run(".")
    doc.save(str(src))

    translate_docx(src, dst, _PreservingMockTranslator(), preserve_runs=True)

    out_runs = list(Document(str(dst)).paragraphs[0].runs)
    assert [r.text for r in out_runs] == ["THE ", "CONTRACT", " IS ", "VOID", "."]
    # Style preserved
    assert out_runs[1].bold is True
    assert out_runs[3].italic is True


def test_run_preserving_falls_back_cleanly(tmp_path) -> None:
    """If the translator drops placeholders, the docx should still
    contain the translation — collapsed into the first run, but never
    garbled."""
    pytest.importorskip("docx")
    from docx import Document

    from nom.translate.formats import translate_docx

    src = tmp_path / "src.docx"
    dst = tmp_path / "dst.docx"
    doc = Document()
    para = doc.add_paragraph("")
    para.add_run("The ")
    para.add_run("contract").bold = True
    para.add_run(".")
    doc.save(str(src))

    translate_docx(src, dst, _DroppingTranslator(), preserve_runs=True)

    out_runs = list(Document(str(dst)).paragraphs[0].runs)
    # All translated text in run 0; runs 1+ cleared.
    assert out_runs[0].text == "THE CONTRACT."
    assert out_runs[1].text == ""
    assert out_runs[2].text == ""


def test_tag_protect_result_dataclass_attributes() -> None:
    result = TagProtectResult(protected=True, run_texts=["a"], fallback_text="a")
    assert result.protected is True
    assert result.run_texts == ["a"]
    assert result.fallback_text == "a"
