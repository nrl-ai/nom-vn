"""Image → DOCX via Tesseract OCR.

Single-image input (.png, .jpg, .jpeg, .tiff, .bmp, .webp) → editable
.docx with the OCR'd text laid out as paragraphs. Vietnamese is the
default OCR language with English as the fallback / co-language;
override via ``ocr_language``.

The OCR text is paragraph-segmented on blank lines (Tesseract emits
double-newline boundaries between paragraphs in --psm 3 default).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING, Any

from nom.convert.base import ConversionStats

if TYPE_CHECKING:
    from PIL import Image as PILImage  # noqa: F401

__all__ = ["image_to_docx"]


_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n+")


def image_to_docx(
    src: Path | str,
    dst: Path | str,
    *,
    ocr_language: str = "vie+eng",
) -> ConversionStats:
    """OCR ``src`` and write the result to ``dst`` as a ``.docx``.

    The destination's parent directory is created if needed; an
    existing file at ``dst`` is overwritten.
    """
    try:
        from PIL import Image
    except ImportError as exc:
        raise ImportError(
            "Image→DOCX requires Pillow. Install with: pip install nom-vn[doc]"
        ) from exc
    try:
        import pytesseract
    except ImportError as exc:
        raise ImportError(
            "Image→DOCX requires pytesseract + Tesseract binary. "
            "Install with: pip install nom-vn[doc] and apt install tesseract-ocr tesseract-ocr-vie"
        ) from exc
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "Image→DOCX requires python-docx. Install with: pip install nom-vn[doc]"
        ) from exc

    src_path = Path(src)
    dst_path = Path(dst)
    if not src_path.exists():
        raise FileNotFoundError(f"image source not found: {src_path}")

    with Image.open(str(src_path)) as img:
        text = pytesseract.image_to_string(img, lang=ocr_language)

    chars_out = _write_docx(Document(), text, dst_path)
    return ConversionStats(
        n_pages=1,
        pages_text_extracted=0,
        pages_ocred=1,
        chars_out=chars_out,
        ocr_language=ocr_language,
    )


def _write_docx(doc: Any, text: str, dst_path: Path) -> int:
    """Write ``text`` as paragraphs into ``doc`` and save to ``dst_path``.

    Returns total character count of the non-empty paragraphs written.
    """
    chars = 0
    paragraphs = _PARAGRAPH_SPLIT.split(text.strip())
    for para in paragraphs:
        cleaned = para.strip()
        if not cleaned:
            continue
        doc.add_paragraph(cleaned)
        chars += len(cleaned)
    if chars == 0:
        # Always write at least one paragraph so the .docx is non-empty.
        doc.add_paragraph("")
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst_path))
    return chars
