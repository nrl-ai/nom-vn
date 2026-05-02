"""PDF → DOCX with a text-layer-first, OCR-fallback strategy.

Each PDF page is processed independently:

1. Try to extract text via ``pdfplumber``. If the page yields more
   than ``min_chars_text_layer`` characters, treat as a "text-layer"
   page — extract the layer directly, no OCR.
2. Otherwise (scanned / image-only page), render the page to a
   bitmap via ``pypdfium2`` at ``ocr_dpi`` and OCR via Tesseract.

The output ``.docx`` is one paragraph per source paragraph (split on
blank lines), with a hard page break between pages so the page
boundaries of the original survive into the editable doc.

Trade-offs:

- Layout drift is real — text-layer extraction yields running text,
  not pixel-positioned blocks. Multi-column PDFs may interleave;
  v0 accepts this. For pixel-fidelity output, route through a
  layout-aware tool (PDFMathTranslate-style); deferred.
- Tables and images pass through as text where possible (pdfplumber's
  flow handles simple tables); complex tables and figures may not
  survive intact.
- Headers / footers / page numbers are usually re-extracted on every
  page; consider de-duplicating in v0.5.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from nom.convert.base import ConversionStats

__all__ = ["pdf_to_docx"]


_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n+")


def pdf_to_docx(
    src: Path | str,
    dst: Path | str,
    *,
    ocr_language: str = "vie+eng",
    ocr_dpi: int = 220,
    min_chars_text_layer: int = 32,
) -> ConversionStats:
    """Convert a PDF to a ``.docx``, using OCR for image-only pages.

    Args:
        src: input PDF.
        dst: output ``.docx``. Parent dir created if missing; overwritten
            if it exists.
        ocr_language: Tesseract language pack (default ``vie+eng``).
        ocr_dpi: render DPI for the OCR fallback. Higher = better
            recognition but slower; 220 is a reasonable middle.
        min_chars_text_layer: a page is considered to have a usable text
            layer when ``pdfplumber`` yields at least this many
            characters. Below the threshold, fall back to OCR.
    """
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError(
            "PDF→DOCX requires pdfplumber. Install with: pip install nom-vn[doc]"
        ) from exc
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise ImportError(
            "PDF→DOCX requires pypdfium2 for the OCR-fallback path. "
            "Install with: pip install nom-vn[doc]"
        ) from exc
    try:
        import pytesseract
    except ImportError as exc:
        raise ImportError(
            "PDF→DOCX OCR fallback requires pytesseract + Tesseract binary. "
            "Install with: pip install nom-vn[doc] and apt install tesseract-ocr tesseract-ocr-vie"
        ) from exc
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
    except ImportError as exc:
        raise ImportError(
            "PDF→DOCX requires python-docx. Install with: pip install nom-vn[doc]"
        ) from exc

    src_path = Path(src)
    dst_path = Path(dst)
    if not src_path.exists():
        raise FileNotFoundError(f"PDF source not found: {src_path}")

    doc = Document()
    n_pages = 0
    pages_text_extracted = 0
    pages_ocred = 0
    chars_out = 0

    pdfium_doc = pdfium.PdfDocument(str(src_path))
    try:
        with pdfplumber.open(str(src_path)) as pdf:
            n_pages = len(pdf.pages)
            for idx, page in enumerate(pdf.pages):
                text = (page.extract_text() or "").strip()
                # pdfplumber returns "(cid:NN)" tokens when the PDF
                # uses subset fonts without a ToUnicode CMap (UDHR is
                # the canonical example). The string is long but not
                # readable — force the OCR fallback for those pages.
                cid_glyph_ratio = (
                    text.count("(cid:") / max(len(text) // 8, 1) if text else 0.0
                )
                usable_text_layer = (
                    len(text) >= min_chars_text_layer and cid_glyph_ratio < 0.05
                )
                if usable_text_layer:
                    pages_text_extracted += 1
                else:
                    text = _ocr_page(
                        pdfium_doc, idx, dpi=ocr_dpi, language=ocr_language, pytesseract=pytesseract
                    )
                    pages_ocred += 1

                chars_out += _write_page(doc, text, is_first_page=(idx == 0), wd_break=WD_BREAK)
    finally:
        pdfium_doc.close()

    if not doc.paragraphs:
        doc.add_paragraph("")  # never write a zero-paragraph .docx

    dst_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst_path))

    return ConversionStats(
        n_pages=n_pages,
        pages_text_extracted=pages_text_extracted,
        pages_ocred=pages_ocred,
        chars_out=chars_out,
        ocr_language=ocr_language,
    )


def _ocr_page(pdfium_doc: Any, page_idx: int, *, dpi: int, language: str, pytesseract: Any) -> str:
    """Render PDF page ``page_idx`` to a PIL image at ``dpi`` and OCR it."""
    page = pdfium_doc[page_idx]
    pil_image = page.render(scale=dpi / 72).to_pil()
    text: str = pytesseract.image_to_string(pil_image, lang=language)
    page.close()
    return text


def _write_page(doc: Any, text: str, *, is_first_page: bool, wd_break: Any) -> int:
    """Append ``text`` to ``doc`` as paragraphs, with a page break before
    each page after the first. Returns characters written."""
    if not is_first_page:
        # Insert a hard page break in front of this page's content.
        doc.add_paragraph().add_run().add_break(wd_break.PAGE)

    chars = 0
    paragraphs = _PARAGRAPH_SPLIT.split(text.strip())
    for para in paragraphs:
        cleaned = para.strip()
        if not cleaned:
            continue
        doc.add_paragraph(cleaned)
        chars += len(cleaned)
    return chars
