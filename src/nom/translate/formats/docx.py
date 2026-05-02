"""Format-preserving .docx translation walker.

Walks the python-docx tree (body, tables, headers / footers across all
sections) and translates each paragraph as a unit.

**v0 strategy: paragraph-level styling.** All runs in a paragraph
(including those nested inside hyperlinks) are concatenated, translated
as one string, written into the first run, with subsequent runs
cleared. The first run's style propagates across the whole paragraph.

Trade-off: mixed-style paragraphs lose intra-paragraph styling. A
sentence with one bold word inside an otherwise-plain paragraph becomes
plain after translation. Paragraph-level styling — headings, list
levels, font, alignment, color, spacing — is fully preserved.

Out of scope for v0 (passes through unchanged):

- Footnotes, endnotes, comments
- Tracked changes, revisions
- Embedded objects (charts, SmartArt, OLE)
- Equations, math content (`<m:oMath>`)
- Image alt-text

A v0.5 path with proportional run-redistribution and a v1 path with
mBERT word-alignment are described in the design doc; both deferred
until user feedback proves they're needed.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any

from nom.translate.base import Translator

if TYPE_CHECKING:
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

__all__ = ["DocxTranslationStats", "translate_docx"]


@dataclass(frozen=True, slots=True)
class DocxTranslationStats:
    """Summary of a single :func:`translate_docx` run.

    ``paragraphs_translated`` counts paragraphs where the translator
    returned successfully. ``paragraphs_skipped`` are empty / whitespace
    paragraphs we passed through. ``paragraphs_failed`` are paragraphs
    where the translator raised — those are left unchanged in the
    output, so the doc remains valid even when the model trips on a
    specific input.
    """

    paragraphs_translated: int
    paragraphs_skipped: int
    paragraphs_failed: int
    chars_in: int
    chars_out: int


def translate_docx(
    src: Path | str,
    dst: Path | str,
    translator: Translator,
    *,
    preserve_runs: bool = False,
) -> DocxTranslationStats:
    """Translate a ``.docx`` file from ``src`` to ``dst``.

    The source file is left untouched. The destination is written
    fresh; any existing file at ``dst`` is overwritten.

    Walks: body paragraphs, body tables (recursively for nested
    tables), and header / footer paragraphs and tables across every
    section (default, first-page, even-page).

    Args:
        src: input ``.docx``.
        dst: output ``.docx`` (parent created if missing; overwritten
            if it exists).
        translator: any :class:`~nom.translate.Translator`.
        preserve_runs: if ``True``, attempt v1 tag-protection — wrap
            run boundaries with ``⟦N⟧`` placeholders, ask the translator
            to preserve them, and redistribute target text into the
            original runs so intra-paragraph styling (mid-sentence bold,
            italic, hyperlinks) survives. Falls back transparently to
            v0 paragraph-level styling per paragraph when the
            placeholders don't round-trip cleanly. Default ``False``
            keeps the proven v0 behaviour.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "Translating .docx requires python-docx. Install with: pip install nom-vn[doc]"
        ) from exc

    src_path = Path(src)
    dst_path = Path(dst)
    if not src_path.exists():
        raise FileNotFoundError(f"docx source not found: {src_path}")

    doc = Document(str(src_path))
    counts = _Counts()
    # lxml-backed XML elements are hashable by underlying C-node address,
    # which is stable across Python wrapper recreation — `id()` on the
    # python-docx wrapper is NOT (proxies get GC'd between iterations,
    # reusing memory). Dedup is needed for merged cells, which expose
    # the same <w:p> via multiple <w:tc> handles.
    seen: set[Any] = set()

    for para in doc.paragraphs:
        _translate_paragraph(para, translator, counts, seen, preserve_runs=preserve_runs)
    for tbl in doc.tables:
        _translate_table(tbl, translator, counts, seen, preserve_runs=preserve_runs)

    for section in doc.sections:
        for region in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if region is None:
                continue
            for para in region.paragraphs:
                _translate_paragraph(para, translator, counts, seen, preserve_runs=preserve_runs)
            for tbl in region.tables:
                _translate_table(tbl, translator, counts, seen, preserve_runs=preserve_runs)

    dst_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst_path))
    return counts.freeze()


# ----------------------------------------------------------------------
# Internals


@dataclass
class _Counts:
    translated: int = 0
    skipped: int = 0
    failed: int = 0
    chars_in: int = 0
    chars_out: int = 0

    def freeze(self) -> DocxTranslationStats:
        return DocxTranslationStats(
            paragraphs_translated=self.translated,
            paragraphs_skipped=self.skipped,
            paragraphs_failed=self.failed,
            chars_in=self.chars_in,
            chars_out=self.chars_out,
        )


def _translate_paragraph(
    paragraph: Paragraph,
    translator: Translator,
    counts: _Counts,
    seen: set[Any],
    *,
    preserve_runs: bool = False,
) -> None:
    element = paragraph._element
    if element in seen:
        return
    seen.add(element)

    runs = _all_runs_in_order(paragraph)
    source = "".join(r.text for r in runs) if runs else ""
    if not source:
        return
    if not source.strip():
        counts.skipped += 1
        return

    counts.chars_in += len(source)

    if preserve_runs and len(runs) > 1:
        from nom.translate._protect import translate_with_tag_protection

        try:
            result = translate_with_tag_protection(
                [r.text for r in runs],
                translator,
            )
        except Exception:
            counts.failed += 1
            return

        if result.protected:
            for run, new_text in zip(runs, result.run_texts, strict=False):
                run.text = new_text
            counts.chars_out += sum(len(t) for t in result.run_texts)
            counts.translated += 1
            return
        # Tag protection round-trip failed — fall through to v0 collapse
        # using the cleaned fallback text the protector returned.
        runs[0].text = result.fallback_text
        for run in runs[1:]:
            run.text = ""
        counts.chars_out += len(result.fallback_text)
        counts.translated += 1
        return

    # v0 collapse — single-run paragraph, or preserve_runs disabled.
    try:
        translated = translator.translate(source)
    except Exception:
        counts.failed += 1
        return

    counts.chars_out += len(translated)
    runs[0].text = translated
    for run in runs[1:]:
        run.text = ""
    counts.translated += 1


def _translate_table(
    table: Table,
    translator: Translator,
    counts: _Counts,
    seen: set[Any],
    *,
    preserve_runs: bool = False,
) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _translate_paragraph(para, translator, counts, seen, preserve_runs=preserve_runs)
            for nested in cell.tables:
                _translate_table(nested, translator, counts, seen, preserve_runs=preserve_runs)


def _all_runs_in_order(paragraph: Paragraph) -> list[Run]:
    """Every visible run in document order — direct children plus runs
    nested inside ``<w:hyperlink>``. python-docx's ``paragraph.runs``
    skips hyperlink-nested runs, which would silently drop translation
    for hyperlinked text (common in legal docs that cite clause
    numbers as links).
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    out: list[Run] = []
    element: Any = paragraph._element
    for child in element.iterchildren():
        tag = child.tag
        if tag == qn("w:r"):
            out.append(Run(child, paragraph))
        elif tag == qn("w:hyperlink"):
            for r_elem in child.findall(qn("w:r")):
                out.append(Run(r_elem, paragraph))
    return out
